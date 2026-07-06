"""
docx_builder.py
~~~~~~~~~~~~~~~
Converts the structured JSON dict produced by llm_service.py into a
professionally formatted .docx end-of-internship report using the
python-docx library.

Layout
------
  Cover page
  Table of Contents (placeholder — Word refreshes on open)
  Abstract
  Acknowledgements
  Numbered chapters & sections
  References
  Appendices (if any)
"""

import io
import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Helpers ───────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    """Fill a table cell background with a hex colour string (e.g. '2E75B6')."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_horizontal_rule(doc: Document, color: str = "2E75B6"):
    """Add a thin coloured horizontal rule paragraph."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p


def _add_page_break(doc: Document):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx_break_type())
    return p


def docx_break_type():
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    return br  # This is used differently — see _add_page_break


def _page_break(doc: Document):
    """Insert a proper page break."""
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    return p


def _set_run_font(run, name: str = "Times New Roman", size_pt: int = 12,
                  bold: bool = False, italic: bool = False,
                  color: RGBColor = None):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def _add_styled_heading(doc: Document, text: str, level: int):
    """Add a heading with custom font styling on top of built-in heading styles."""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.name = "Times New Roman"
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)  # Dark blue
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)  # Medium blue
        else:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    return heading


def _add_body_paragraph(doc: Document, text: str):
    """Add a justified body-text paragraph with proper academic spacing."""
    if not text.strip():
        return
    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.first_line_indent = Inches(0.5)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        run = p.add_run(block)
        _set_run_font(run)


# ── Cover page ────────────────────────────────────────────────────────────────

def _build_cover_page(doc: Document, report_data: dict, metadata: dict):
    """Create a professional cover page."""
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Top accent bar via paragraph border
    _add_horizontal_rule(doc, "1F497D")

    # Spacer
    for _ in range(3):
        doc.add_paragraph()

    # Institution name
    inst = metadata.get("institution_name") or "University / Institution"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(inst.upper())
    _set_run_font(run, size_pt=14, bold=True, color=RGBColor(0x1F, 0x49, 0x7D))

    # Programme
    prog = metadata.get("programme") or ""
    if prog:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(prog)
        _set_run_font(run2, size_pt=12, italic=True)

    doc.add_paragraph()
    _add_horizontal_rule(doc, "2E75B6")
    doc.add_paragraph()

    # Report title
    title = report_data.get("title", "End-of-Internship Report")
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(title.upper())
    _set_run_font(run_title, size_pt=20, bold=True, color=RGBColor(0x1F, 0x49, 0x7D))

    doc.add_paragraph()
    _add_horizontal_rule(doc, "2E75B6")

    # Metadata table
    for _ in range(2):
        doc.add_paragraph()

    fields = [
        ("Submitted by",   metadata.get("intern_name") or "[Intern Name]"),
        ("Host Organisation", metadata.get("company_name") or "[Company Name]"),
        ("Department",     metadata.get("department") or "[Department]"),
        ("Supervisor",     metadata.get("supervisor_name") or "[Supervisor Name]"),
        ("Duration",       metadata.get("internship_duration") or "[Duration]"),
        ("Date",           datetime.date.today().strftime("%B %Y")),
    ]

    table = doc.add_table(rows=len(fields), cols=2)
    table.style = "Table Grid"
    col_widths = [Inches(2.0), Inches(4.0)]

    for i, (label, value) in enumerate(fields):
        row = table.rows[i]
        # Label cell
        lc = row.cells[0]
        lc.width = col_widths[0]
        _set_cell_bg(lc, "1F497D")
        lp = lc.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        lr = lp.add_run(label)
        _set_run_font(lr, size_pt=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        # Value cell
        vc = row.cells[1]
        vc.width = col_widths[1]
        _set_cell_bg(vc, "EBF3FB")
        vp = vc.paragraphs[0]
        vr = vp.add_run(value)
        _set_run_font(vr, size_pt=11)

    for _ in range(3):
        doc.add_paragraph()
    _add_horizontal_rule(doc, "1F497D")

    _page_break(doc)


# ── Table of Contents placeholder ────────────────────────────────────────────

def _build_toc(doc: Document):
    """Insert a TOC placeholder that Word refreshes on open."""
    _add_styled_heading(doc, "Table of Contents", level=1)

    p = doc.add_paragraph()
    run = p.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar)

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instrText)

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    run._r.append(fldChar2)

    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar3)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_run = p2.add_run(
        "(Right-click and select 'Update Field' to refresh this table of contents)"
    )
    _set_run_font(note_run, size_pt=9, italic=True, color=RGBColor(0x80, 0x80, 0x80))

    _page_break(doc)


# ── Main builder ─────────────────────────────────────────────────────────────

def build_docx(report_data: dict, metadata: dict) -> bytes:
    """
    Build the complete .docx document from the LLM-generated structure.

    Parameters
    ----------
    report_data : dict
        Parsed JSON from the LLM (matches the schema in llm_service.py).
    metadata : dict
        Original request metadata (intern_name, company_name, etc.).

    Returns
    -------
    bytes
        The raw .docx file bytes ready to be streamed to the client.
    """
    doc = Document()

    # ── Global styles ────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # ── Cover page ───────────────────────────────────────────────────────────
    _build_cover_page(doc, report_data, metadata)

    # ── Table of Contents ────────────────────────────────────────────────────
    _build_toc(doc)

    # ── Abstract ─────────────────────────────────────────────────────────────
    abstract = report_data.get("abstract", "")
    if abstract:
        _add_styled_heading(doc, "Abstract", level=1)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.right_indent = Inches(0.5)
        run = p.add_run(abstract)
        _set_run_font(run, italic=True)
        _page_break(doc)

    # ── Acknowledgements ─────────────────────────────────────────────────────
    ack = report_data.get("acknowledgements", "")
    if ack:
        _add_styled_heading(doc, "Acknowledgements", level=1)
        _add_body_paragraph(doc, ack)
        _page_break(doc)

    # ── Body sections ────────────────────────────────────────────────────────
    chapter_counter = 0
    sections = report_data.get("sections", [])

    for sec in sections:
        level = int(sec.get("level", 1))
        heading_text = sec.get("heading", "Section")
        body = sec.get("body", "")

        if level == 1:
            chapter_counter += 1
            numbered = f"{chapter_counter}. {heading_text}"
            _add_styled_heading(doc, numbered, level=1)
        elif level == 2:
            _add_styled_heading(doc, heading_text, level=2)
        else:
            _add_styled_heading(doc, heading_text, level=3)

        _add_body_paragraph(doc, body)

        # Page break after each top-level chapter
        if level == 1:
            _page_break(doc)

    # ── Conclusion ───────────────────────────────────────────────────────────
    conclusion = report_data.get("conclusion", "")
    if conclusion:
        chapter_counter += 1
        _add_styled_heading(doc, f"{chapter_counter}. Conclusion", level=1)
        _add_body_paragraph(doc, conclusion)
        _page_break(doc)

    # ── Recommendations ──────────────────────────────────────────────────────
    recommendations = report_data.get("recommendations", [])
    if recommendations:
        chapter_counter += 1
        _add_styled_heading(doc, f"{chapter_counter}. Recommendations", level=1)
        for i, rec in enumerate(recommendations, 1):
            p = doc.add_paragraph(style="List Number")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(rec)
            _set_run_font(run)
        _page_break(doc)

    # ── References ───────────────────────────────────────────────────────────
    references = report_data.get("references", [])
    if references:
        _add_styled_heading(doc, "References", level=1)
        for ref in references:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(ref)
            _set_run_font(run, size_pt=11)

    # ── Serialise to bytes ───────────────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
